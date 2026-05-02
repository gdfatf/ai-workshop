from __future__ import annotations

from pathlib import Path
import os
from typing import Optional, List, Any, Dict

import streamlit as st
from dotenv import load_dotenv

# ========= LangChain / Providers =========
from langchain_anthropic import ChatAnthropic
from langchain_openai import OpenAIEmbeddings

# ========= LangChain core =========
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.documents import Document

# ========= RAG / Vector DB =========
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import Chroma

# ========= DOCX loader =========
from docx import Document as DocxDocument

# ========= Supabase =========
from supabase import create_client, Client


# =========================
# Streamlit 基础设置
# =========================
st.set_page_config(page_title="AI Workbench · 4-Agent 短视频创作系统", layout="wide")
st.title("🎬 AI Workbench · 4-Agent 短视频全链路创作系统")


# =========================
# Env/Secrets
# =========================
load_dotenv()


def sget(key: str, default: Optional[str] = None) -> Optional[str]:
    try:
        if hasattr(st, "secrets") and key in st.secrets:
            return str(st.secrets[key])
    except FileNotFoundError:
        pass
    return os.getenv(key, default)


ANTHROPIC_API_KEY = sget("ANTHROPIC_API_KEY")
LLM_MODEL_DEFAULT = sget("LLM_MODEL", "claude-opus-4-7")
OPENAI_API_KEY = sget("OPENAI_API_KEY")
APP_PASSWORD = sget("APP_PASSWORD", "")

OPENAI_EMBED_MODEL_DEFAULT = sget("OPENAI_EMBED_MODEL", "text-embedding-3-large")

SUPABASE_URL = sget("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = sget("SUPABASE_SERVICE_ROLE_KEY")

if not ANTHROPIC_API_KEY:
    st.error("缺少 ANTHROPIC_API_KEY：请在 Streamlit Secrets 或本地 .env 中配置。")
    st.stop()

if not OPENAI_API_KEY:
    st.error("缺少 OPENAI_API_KEY：请在 Streamlit Secrets 或本地 .env 中配置（用于 Embedding / RAG）。")
    st.stop()

if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
    st.error("缺少 SUPABASE_URL / SUPABASE_SERVICE_ROLE_KEY：请在 Streamlit Secrets 配置（用于记忆存储）。")
    st.stop()


@st.cache_resource(show_spinner=False)
def get_supabase() -> Client:
    return create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)


sb = get_supabase()


# =========================
# 简单密码门（可选）
# =========================
if APP_PASSWORD:
    if "authed" not in st.session_state:
        st.session_state.authed = False

    if not st.session_state.authed:
        st.subheader("🔒 请输入访问密码")
        pwd = st.text_input("Password", type="password", key="pwd_input")
        if st.button("进入", key="pwd_btn"):
            if pwd == APP_PASSWORD:
                st.session_state.authed = True
                st.rerun()
            else:
                st.error("密码不正确")
        st.stop()


# =========================
# 路径与目录
# =========================
PROMPTS_DIR = Path("agents") / "prompts"
KB_DIR = Path("kb")
DB_DIR = Path("chroma_db")

PROMPTS_DIR.mkdir(parents=True, exist_ok=True)
KB_DIR.mkdir(parents=True, exist_ok=True)
DB_DIR.mkdir(parents=True, exist_ok=True)

# =====================================================
# 4-Agent 定义
# =====================================================
AGENTS = {
    "A｜账号定位师 · 战略罗盘": "agent_a.txt",
    "B｜选题架构师 · 爆款引擎": "agent_b.txt",
    "C｜文案炼金师 · 结构外科": "agent_c.txt",
    "D｜影像导演 · 情绪建筑师": "agent_d.txt",
}

AGENT_DESCRIPTIONS = {
    "A｜账号定位师 · 战略罗盘": "定位诊断 → 赛道验证 → 变现设计 → 人设锻造",
    "B｜选题架构师 · 爆款引擎": "灵感捕获 → 选题打磨 → 流量漏斗 → 架构选择",
    "C｜文案炼金师 · 结构外科": "结构体检 → 钩子手术 → 骨架重排 → 刺点植入 → 语言雕琢",
    "D｜影像导演 · 情绪建筑师": "情感团块 → 风格定调 → 分镜脚本 → 表演指导",
}


# =========================
# Supabase Memory
# =========================
def db_load_chat(user_id: str, agent_id: str) -> List[Dict[str, str]]:
    try:
        resp = (
            sb.table("workbench_memory")
            .select("messages")
            .eq("user_id", user_id)
            .eq("agent_id", agent_id)
            .execute()
        )
        data = resp.data
        if not data:
            return []
        row = data[0] if isinstance(data, list) else data
        msgs = row.get("messages", [])
        return msgs if isinstance(msgs, list) else []
    except Exception:
        return []


def db_save_chat(user_id: str, agent_id: str, messages: List[Dict[str, str]]) -> None:
    try:
        sb.table("workbench_memory").upsert(
            {
                "user_id": user_id,
                "agent_id": agent_id,
                "messages": messages,
            }
        ).execute()
    except Exception:
        pass


def db_clear_chat(user_id: str, agent_id: str) -> None:
    db_save_chat(user_id, agent_id, [])


def lc_to_json(messages: List[Any]) -> List[Dict[str, str]]:
    out: List[Dict[str, str]] = []
    for m in messages:
        if isinstance(m, HumanMessage):
            out.append({"role": "user", "content": m.content})
        elif isinstance(m, AIMessage):
            out.append({"role": "assistant", "content": m.content})
        elif isinstance(m, SystemMessage):
            out.append({"role": "system", "content": m.content})
    return out


def json_to_lc(items: List[Dict[str, str]]) -> List[Any]:
    out: List[Any] = []
    for it in items:
        role = it.get("role")
        content = it.get("content", "")
        if role == "user":
            out.append(HumanMessage(content=content))
        elif role == "assistant":
            out.append(AIMessage(content=content))
        elif role == "system":
            out.append(SystemMessage(content=content))
    return out


# =========================
# 工具函数
# =========================
def load_prompt(filename: str) -> str:
    path = PROMPTS_DIR / filename
    if not path.exists():
        return "You are a helpful assistant."
    txt = path.read_text(encoding="utf-8").strip()
    return txt or "You are a helpful assistant."


def load_docx_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(parts)


def load_kb_documents(agent_id: str) -> List[Document]:
    folder = KB_DIR / agent_id
    folder.mkdir(parents=True, exist_ok=True)

    docs: List[Document] = []
    for p in folder.rglob("*"):
        if p.is_dir():
            continue
        suf = p.suffix.lower()

        if suf in (".txt", ".md"):
            docs.extend(TextLoader(str(p), encoding="utf-8").load())

        elif suf == ".docx":
            text = load_docx_text(p)
            if text.strip():
                docs.append(Document(page_content=text, metadata={"source": str(p)}))

    return docs


def extract_text(resp) -> str:
    content = getattr(resp, "content", resp)

    if isinstance(content, str):
        return content

    if isinstance(content, list):
        out = []
        for block in content:
            if hasattr(block, "text"):
                out.append(str(getattr(block, "text", "")))
            elif isinstance(block, dict):
                out.append(str(block.get("text", "")))
            else:
                out.append(str(block))
        return "".join(out)

    if hasattr(resp, "text"):
        return str(getattr(resp, "text", ""))

    return str(content)


def estimate_msg_len(msg: Any) -> int:
    try:
        return len(getattr(msg, "content", "") or "")
    except Exception:
        return 0


def trim_chat_history_by_chars(
    chat: List[Any],
    max_chars: int = 18000,
    min_keep_turns: int = 12,
) -> List[Any]:
    """
    不按固定轮数硬裁剪，而是按字符预算保留更多上下文。
    同时至少保留最近 min_keep_turns 轮（约 2*turns 条消息）。
    """
    if not chat:
        return []

    min_keep_msgs = min(len(chat), min_keep_turns * 2)
    guaranteed = chat[-min_keep_msgs:]

    # 如果保底部分已经很长，也至少直接返回保底
    total_guaranteed = sum(estimate_msg_len(m) for m in guaranteed)
    if total_guaranteed >= max_chars:
        return guaranteed

    selected = guaranteed[:]
    current_chars = total_guaranteed

    older = chat[:-min_keep_msgs]
    older_reversed = list(reversed(older))

    for msg in older_reversed:
        msg_len = estimate_msg_len(msg)
        if current_chars + msg_len > max_chars:
            break
        selected.insert(0, msg)
        current_chars += msg_len

    return selected


def build_llm(model_name: str, temperature: float, long_mode: bool) -> ChatAnthropic:
    llm_kwargs = {
        "model": model_name,
        "api_key": ANTHROPIC_API_KEY,
        "max_tokens": 5000 if long_mode else 3200,
        "timeout": 300 if long_mode else 240,
        "max_retries": 1,
        "streaming": True,
    }

    if "4-7" not in model_name:
        llm_kwargs["temperature"] = temperature

    return ChatAnthropic(**llm_kwargs)

def build_embeddings(model_name: str) -> OpenAIEmbeddings:
    return OpenAIEmbeddings(
        model=model_name,
        api_key=OPENAI_API_KEY,
    )


@st.cache_resource(show_spinner=False)
def get_vectorstore(agent_id: str, embed_model: str):
    embeddings = build_embeddings(embed_model)

    persist_dir = DB_DIR / agent_id
    persist_dir.mkdir(parents=True, exist_ok=True)

    safe_model = embed_model.replace("/", "_").replace(":", "_")
    collection_name = f"kb_{agent_id}__{safe_model}"

    vs = Chroma(
        collection_name=collection_name,
        embedding_function=embeddings,
        persist_directory=str(persist_dir),
    )

    try:
        existing = vs._collection.count()
    except Exception:
        existing = 0

    if existing == 0:
        raw_docs = load_kb_documents(agent_id)
        if raw_docs:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=1200,
                chunk_overlap=180,
                separators=["\n## ", "\n### ", "\n#### ", "\n---", "\n\n", "\n", " "],
            )
            chunks = splitter.split_documents(raw_docs)
            vs.add_documents(chunks)

    return vs


def retrieve_context(
    agent_id: str,
    query: str,
    k: int,
    embed_model: str,
    max_chars: int = 10000,
    per_doc_max_chars: int = 1800,
) -> str:
    vs = get_vectorstore(agent_id, embed_model=embed_model)
    docs = vs.similarity_search(query, k=k)
    if not docs:
        return ""

    blocks = []
    total = 0

    for i, d in enumerate(docs, 1):
        src = d.metadata.get("source", "kb")
        text = d.page_content.strip()
        text = text[:per_doc_max_chars]

        block = f"[片段{i} | 来源: {src}]\n{text}"
        block_len = len(block)

        if total + block_len > max_chars:
            break

        blocks.append(block)
        total += block_len

    return "\n\n".join(blocks)


# =========================
# Sidebar UI
# =========================
with st.sidebar:
    st.header("⚙️ 设置")

    st.subheader("用户身份（用于记忆隔离）")
    user_id = st.text_input("你的用户ID（每人固定一个）", key="user_id_input")

    if not user_id or not user_id.strip():
        st.warning("请输入用户ID（例如：姓名拼音/代号），用于保存个人记忆。")
        st.stop()
    user_id = user_id.strip()

    st.divider()

    agent_name = st.selectbox("选择 Agent", list(AGENTS.keys()), key="agent_select")

    desc = AGENT_DESCRIPTIONS.get(agent_name, "")
    if desc:
        st.caption(f"📋 {desc}")

    st.divider()

    claude_candidates = [
        "claude-opus-4-7",
        "claude-sonnet-4-7",
        "claude-opus-4-6",
        "claude-sonnet-4-6",
        "claude-haiku-4-5",
    ]

    claude_default = (
        LLM_MODEL_DEFAULT
        if LLM_MODEL_DEFAULT in claude_candidates
        else claude_candidates[0]
    )

    model_name = st.selectbox(
        "LLM（Claude）",
        claude_candidates,
        index=claude_candidates.index(claude_default),
        key="llm_select",
    )

    embed_candidates = [
        "text-embedding-3-large",
        "text-embedding-3-small",
    ]

    embed_default = (
        OPENAI_EMBED_MODEL_DEFAULT
        if OPENAI_EMBED_MODEL_DEFAULT in embed_candidates
        else embed_candidates[0]
    )

    embed_model = st.selectbox(
        "Embedding（OpenAI）",
        embed_candidates,
        index=embed_candidates.index(embed_default),
        key="embed_select",
    )

    temperature = st.slider("temperature", 0.0, 1.0, 0.3, 0.05, key="temp_slider")

    use_rag = st.toggle("启用 RAG（从 KB 检索）", value=True, key="rag_toggle")
    topk = st.slider("检索 TopK", 1, 8, 4, 1, key="topk_slider")

    long_mode = st.toggle("长文/深度模式", value=True, key="long_mode")

    if st.button("🗑️ 清空当前 Agent 对话", key="clear_chat_btn"):
        agent_file_tmp = AGENTS[agent_name]
        agent_id_tmp = agent_file_tmp.replace(".txt", "")
        db_clear_chat(user_id, agent_id_tmp)
        st.session_state.pop(f"chat::{user_id}::{agent_id_tmp}", None)
        st.rerun()

    st.divider()
    st.caption("📂 Prompt → agents/prompts/agent_a~d.txt")
    st.caption("📂 KB → kb/agent_a~d/（.txt .md .docx）")

    st.divider()
    st.caption("🔄 创作流程：A定位 → B选题 → C文案 → D分镜")


# =========================
# 主流程
# =========================
agent_file = AGENTS[agent_name]
system_prompt = load_prompt(agent_file)
agent_id = agent_file.replace(".txt", "")

llm = build_llm(model_name=model_name, temperature=temperature, long_mode=long_mode)

chat_key = f"chat::{user_id}::{agent_id}"
if chat_key not in st.session_state:
    raw = db_load_chat(user_id, agent_id)
    st.session_state[chat_key] = json_to_lc(raw)

chat = st.session_state[chat_key]

with st.expander("查看当前 Agent 的 System Prompt（只读）", expanded=False):
    st.code(system_prompt)

for msg in chat:
    if isinstance(msg, HumanMessage):
        with st.chat_message("user"):
            st.markdown(msg.content)
    else:
        with st.chat_message("assistant"):
            st.markdown(msg.content)

user_text = st.chat_input(f"正在使用：{agent_name}（可粘贴长文本）")

if user_text:
    chat.append(HumanMessage(content=user_text))
    db_save_chat(user_id, agent_id, lc_to_json(chat))

    with st.chat_message("user"):
        st.markdown(user_text)

    rag_context = ""
    if use_rag:
        try:
            rag_context = retrieve_context(
                agent_id=agent_id,
                query=user_text,
                k=topk,
                embed_model=embed_model,
                max_chars=7000 if long_mode else 5000,
                per_doc_max_chars=1800 if long_mode else 1400,
            )
        except Exception as e:
            st.error(f"RAG / Embedding 初始化失败：{e}")
            st.stop()

    recent_chat = trim_chat_history_by_chars(
        chat=chat,
        max_chars=26000 if long_mode else 18000,
        min_keep_turns=16 if long_mode else 12,
    )

    messages: List[Any] = [SystemMessage(content=system_prompt)]

    if rag_context:
        messages.append(
            HumanMessage(
                content=(
                    "以下是与当前问题相关的知识库片段，请优先基于这些内容回答。"
                    "如果引用了片段，请在回答中标注来源片段编号。"
                    "当知识库片段与历史对话冲突时，以当前问题和更直接相关的片段为准。\n\n"
                    f"{rag_context}"
                )
            )
        )

    messages.extend(recent_chat)

    with st.chat_message("assistant"):
        with st.spinner("思考中…"):
            chunks: List[str] = []
            placeholder = st.empty()

            try:
                for chunk in llm.stream(messages):
                    piece = extract_text(chunk)
                    if piece:
                        chunks.append(piece)
                        placeholder.markdown("".join(chunks))
            except Exception as e:
                st.error(f"Claude 调用超时/失败：{e}")
                st.stop()

            answer = "".join(chunks).strip()

            if not answer:
                answer = "模型未返回可解析内容，请重试。"
                placeholder.markdown(answer)

    chat.append(AIMessage(content=answer))
    db_save_chat(user_id, agent_id, lc_to_json(chat))


with st.expander("🧪 调试面板", expanded=False):
    st.write("user_id：", user_id)
    st.write("Agent：", agent_name)
    st.write("agent_id：", agent_id)
    st.write("LLM：", model_name)
    st.write("Embedding：", embed_model)
    st.write("RAG：", use_rag, "TopK=", topk)
    st.write("Long mode：", long_mode)
    st.write("KB path：", str(KB_DIR / agent_id))
    st.write("当前总历史条数：", len(chat))
    st.write(
        "实际发送历史条数：",
        len(trim_chat_history_by_chars(chat, max_chars=26000 if long_mode else 18000, min_keep_turns=16 if long_mode else 12)),
    )
